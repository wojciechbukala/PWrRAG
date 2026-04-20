import json
import warnings
import hashlib
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import io
from pypdf import PdfReader
import docx
import argparse

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from Models import set_vectorstore, add_to_vectorstore, verify_vectorstore, reset_vectorstore


### CONSTANTS
MIN_PAGE_CHARS = 50             # documents with less than MIN_PAGE_CHARS characters are ignored
TIMEOUT = 10                    # timeout for getting request (in seconds)
CHUNK_SIZE = 1500               # number of characters in one chunk
CHUNK_OVERLAP = 200             # overlap window of chunking
BATCH_SIZE = 800                # batch of chunks to add to vectorstore
MAX_CHUNKS_PER_SOURCE = 2500    # maxiumu number of chunks for one source, prevents one source domination
MAX_FILES_PER_SOURCE = 200      # maximum number of files to download from one source

# Initialize a global session for connection pooling
session = requests.Session()
session.headers.update({
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
})


### KEYWORDS TUPLES
KEYWORDS_PWR = ('uczelnia/informacje-ogolne/statut', 'uczelnia/informacje-ogolne/strategia')
KEYWORDS_DEP = ('studenci/programy-studiow', 'studenci/dokumenty', 'studenci/dokumenty-do-pobrania',
             'studenci/indywidualna-organizacja-studiow', 'studenci/uznanie-efektow-uczenia-sie',
             'studenci/dydaktyka/programy-studiow-programy-ksztalcenia', 'studenci/wzory-podan-i-zalacznikow',
             'studenci/karty-przedmiotow', 'studenci/regulamin-studiow', 'studenci/zarzadzenia-dziekana')

KEYWORDS_BIB = ('o-nas/regulamin-udostepniania-zasady-korzystania')
KEYWORDS_CKU = ('studia-podyplomowe/regulamin')
KEYWORDS_BIZ = ('')
KEYWORDS_NAU = ('zadania/ewaluacja-nauki')
KEYWORDS_PROJ = ('dla-wnioskujacych/ogolne-zasady-przygotowania-projektow',
                 'dla-realizujacych/ogolne-zasady-realizacji-projektow',
                 'dla-realizujacych/przewodnik-teta-edu--obszar-op')
KEYWORDS_CRM = ('studenci/program-erasmus-plus/erasmus/erasmus-ka171')
KEYWORDS_BK  = ('')
KEYWORDS_CDD = ('')
KEYWORDS_RJK = ('uszzjk', 'procesy', 'akredytacje', 'polityka-jakosci')
KEYWORDS_KONF = ('')
KEYWORDS_SAM = ('o-samorzadzie/dokumenty/regulaminy-samorzadu',
                'o-samorzadzie/dokumenty/wzory-dokumentow',
                'dla-studenta/wsparcie-dydaktyczne/regulaminy-i-twoje-prawa/regulaminy')
KEYWORDS_REK = ('rekrutacja/akty-prawne/', 'rekrutacja/dokumenty-rekrutacyjne/')
KEYWORDS_SZD = ('rekrutacja/zasady-rekrutacji', 'rekrutacja/dokumenty-rekrutacyjne',
                'doktoranci/ocena-srodokresowa/dokumenty-do-pobrania', 'promotorzy/materialy-dla-promotorow')

KEYWORDS_MAP = {"KEYWORDS_PWR": KEYWORDS_PWR,
                "KEYWORDS_DEP": KEYWORDS_DEP,
                "KEYWORDS_BIB": KEYWORDS_BIB,
                "KEYWORDS_CKU": KEYWORDS_CKU,
                "KEYWORDS_BIZ": KEYWORDS_BIZ,
                "KEYWORDS_NAU": KEYWORDS_NAU,
                "KEYWORDS_PROJ": KEYWORDS_PROJ,
                "KEYWORDS_CRM": KEYWORDS_CRM,
                "KEYWORDS_BK": KEYWORDS_BK,
                "KEYWORDS_CDD": KEYWORDS_CDD,
                "KEYWORDS_RJK": KEYWORDS_RJK,
                "KEYWORDS_KONF": KEYWORDS_KONF,
                "KEYWORDS_SAM": KEYWORDS_SAM,
                "KEYWORDS_REK": KEYWORDS_REK,
                "KEYWORDS_SZD": KEYWORDS_SZD}


def load_pdf_from_bytes(url):
    """Load pdf files from website and return Document object"""
    response = session.get(url, timeout=TIMEOUT)
    response.raise_for_status()

    # check if content is in PDF format
    content_type = response.headers.get('Content-Type', '')
    if 'html' in content_type or not response.content.lstrip().startswith(b'%PDF'):
        raise ValueError(f"URL did not return a PDF (Content-Type: {content_type})")

    # convert to bytes package
    file_content = io.BytesIO(response.content)
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        reader = PdfReader(file_content, strict=False)

    # build table of langchain Document objects
    docs = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ''
        except Exception:
            text = ''
        if len(text.strip()) >= MIN_PAGE_CHARS:
            docs.append(Document(
                page_content=text,
                metadata={"source": url, "page": i+1}
            ))

    if not docs:
        raise ValueError("No extractable text found (likely a scanned image PDF)")

    # return pdf in the form of Document list
    return docs

def load_docx_from_bytes(url):
    """Load docx files from website and return Document object"""
    response = session.get(url, timeout=TIMEOUT)
    response.raise_for_status()

    file_content = io.BytesIO(response.content)

    document = docx.Document(file_content)
    full_text = "\n".join(para.text for para in document.paragraphs).strip()

    if len(full_text) < MIN_PAGE_CHARS:
        raise ValueError("No extractable text found (empty or image-only DOCX)")

    return [Document(page_content=full_text, metadata={"source": url, "format": "docx"})]

def make_chunks(url, file_type,
          split_chunk_size=CHUNK_SIZE,
          split_chunk_overlap=CHUNK_OVERLAP,
          split_separators=["\n\n", "\n", " ", ""],
          verbose=False):
    """Split retrieved docuents into chunks with RecursiveCharacterTextSplitter"""
    
    # recursive splitter - takes 1500 chars or one paragraph or one line
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=split_chunk_size,
        chunk_overlap=split_chunk_overlap,
        separators=split_separators
    )

    try:
        if file_type == '.pdf':
            docs = load_pdf_from_bytes(url)
        elif file_type == '.docx':
            docs = load_docx_from_bytes(url)
        else:
            if verbose:
                print(f"[ERROR]: Unknown file type: {file_type}")
            return []
        
        chunks = splitter.split_documents(docs)
        if verbose:
            print(f"Created {len(chunks)} chunks from {url}")
        return chunks
    except Exception as e:
        if verbose:
            print(f"[ERROR]: Error chunking {url}: {str(e)}")
        return []


def debug_crawler(base_url, keywords=KEYWORDS_DEP):
    """Dry-run crawl: prints discovered subpages and files without downloading or embedding."""
    crawled = set()
    to_crawl = {base_url}
    subpages = []
    files = []

    print(f"\n=== DEBUG CRAWL: {base_url} ===")
    print(f"Keywords: {keywords}\n")

    while to_crawl:
        url = to_crawl.pop()
        if url in crawled:
            continue
        try:
            response = session.get(url, timeout=TIMEOUT)
            crawled.add(url)

            content_type = response.headers.get('Content-Type', '')
            parser = 'xml' if 'xml' in content_type else 'html.parser'
            bs = BeautifulSoup(response.text, parser)

            for link in bs.find_all('a', href=True):
                potential_url = urljoin(url, link['href'])

                if urlparse(potential_url).netloc != urlparse(base_url).netloc:
                    continue

                ext = None
                if potential_url.lower().endswith('.pdf'): ext = '.pdf'
                elif potential_url.lower().endswith('.docx'): ext = '.docx'

                if ext and potential_url not in files:
                    files.append(potential_url)
                    if len(files) >= MAX_FILES_PER_SOURCE:
                        print(f"Debug: Reached MAX_FILES_PER_SOURCE ({MAX_FILES_PER_SOURCE})")
                        to_crawl = set()
                        break
                    continue

                if potential_url not in crawled:
                    matched = [kw for kw in keywords if kw in potential_url.lower()]
                    if matched:
                        subpages.append((potential_url, matched))
                        to_crawl.add(potential_url)
        except Exception:
            pass

    print(f"--- Subpages found ({len(subpages)}) ---")
    for page_url, matched_kw in sorted(subpages):
        print(f"  [{', '.join(matched_kw)}] {page_url}")

    print(f"\n--- Files found ({len(files)}) ---")
    for f in sorted(files):
        print(f"  {f}")

    print(f"\nTotal: {len(subpages)} subpages, {len(files)} files, {len(crawled)} pages crawled.")


def crawler(base_url, keywords = KEYWORDS_DEP, embed = False, verbose = False):
    """Crawl website to find formal documents based on the defined keywords"""
    crawled = set()
    to_crawl = {base_url}
    files = set()

    chunks = []

    # check all to_crawl urls
    while to_crawl:
        if len(files) >= MAX_FILES_PER_SOURCE:
            if verbose:
                print(f"  Reached threshold of {MAX_FILES_PER_SOURCE} files for {base_url}. Stopping crawl.")
            break

        url = to_crawl.pop() # take one url
        if url in crawled:
            continue
    
        try:
            if verbose:
                print(f'  [{len(crawled)+1}] crawling ({len(to_crawl)} in queue, {len(files)} files found): {url}', flush=True)
            response = session.get(url, timeout=5)
            crawled.add(url) # add url to crawled

            content_type = response.headers.get('Content-Type', '')
            parser = 'xml' if 'xml' in content_type else 'html.parser'
            bs = BeautifulSoup(response.text, parser)

            # look for links
            for link in bs.find_all('a', href=True):
                href = link['href']
                potential_url = urljoin(url, href)

                # ignore external url
                if urlparse(potential_url).netloc != urlparse(base_url).netloc:
                    continue

                # check extention
                extention = None
                if potential_url.lower().endswith('.pdf'): extention = '.pdf'
                elif potential_url.lower().endswith('.docx'): extention = '.docx'

                # look for files
                if extention and potential_url not in files:
                    files.add(potential_url)
                    if verbose: 
                        print(f'Found file: {potential_url}')
                    if embed: 
                        new_chunks = make_chunks(potential_url, extention, verbose=verbose)
                        chunks.extend(new_chunks)
                    
                    if len(files) >= MAX_FILES_PER_SOURCE:
                        break
                    continue

                # look for new paths
                if potential_url not in crawled:
                    if any(word in potential_url.lower() for word in keywords):
                        to_crawl.add(potential_url)
        except Exception:
            pass

    return chunks


def load_sources(sources_file='sources.json', verbose=False):
    """Load sources JSON."""
    try:
        with open(sources_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data.get('sources', [])
    except Exception as e:
        if verbose:
            print(f"[ERROR]: Error loading: {sources_file}: {str(e)}")
        return []
    
def build_vectorstore(sources_json_path: str, reset_vstore=True, verbose=False):

    # load source urls from json
    sources = load_sources(sources_json_path)
    
    if not sources:
        print("Sources not found")
        exit(1)

    # reset vectorstore and build new one
    if reset_vstore:
        reset_vectorstore()

    # define buffer for chunks and hashes to detect duplicated chunks
    buffer = []
    seen_hashes = set()
    total_chunks = 0
    vectorstore_initialized = False

    # iterate through sources
    for i, source in enumerate(sources, 1):
        url = source.get('url')
        dept = source.get('department', 'Unknown')
        keywords = source.get('keywords')
        keywords_tuple = KEYWORDS_MAP[keywords]

        # skip iteration if url not found
        if not url:
            continue
    
        print(f"[{i}/{len(sources)}] Crawling: {dept} ({url})")
        try:
            chunks = crawler(url, keywords=keywords_tuple, embed=True, verbose=verbose)

            if chunks:
                # cut array if exceeded max chunks value, prevention of one source domination
                if len(chunks) > MAX_CHUNKS_PER_SOURCE:
                    print(f"Truncating {len(chunks)} → {MAX_CHUNKS_PER_SOURCE} chunks for {url}")
                    chunks = chunks[:MAX_CHUNKS_PER_SOURCE]

                # check if chunks are not duplicated
                unique_chunks = []
                for chunk in chunks:
                    h = hashlib.md5(chunk.page_content.encode()).hexdigest()
                    if h not in seen_hashes:
                        seen_hashes.add(h)
                        unique_chunks.append(chunk)
                
                skipped = len(chunks) - len(unique_chunks)
                if skipped and verbose:
                    print(f"Skipped {skipped} duplicate chunks from {url}")

                buffer.extend(unique_chunks)
                total_chunks += len(unique_chunks)
                print(f"Found {len(unique_chunks)} chunks (total: {total_chunks})\n")
            else:
                print(f"[WARNING]: No chunks found for {url}\n")

        except Exception as e:
            print(f"[Error]: Crawling error: {dept}: {str(e)}\n")

        # if batch is full - save batch chunks to vectorstore and realise buffer
        if len(buffer) >= BATCH_SIZE:
            print(f"Saving batch of {len(buffer)} chunks...")
            if not vectorstore_initialized:
                set_vectorstore(buffer)
                vectorstore_initialized = True
            else:
                add_to_vectorstore(buffer)
            buffer.clear()

    # save all remains to vectorstore
    if buffer:
        print(f"Saving batch of {len(buffer)} chunks...")
        if not vectorstore_initialized:
            set_vectorstore(buffer)
            vectorstore_initialized = True
        else:
            add_to_vectorstore(buffer)

    # verify vectorstore
    if vectorstore_initialized:
        verify_vectorstore()
    else:
        print("WARNING: No chunks to process.")


if __name__ == '__main__':

    arg_parser = argparse.ArgumentParser(
        description='Web Scraper for PWr formal documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""Use cases:
        python3 Scraper.py     # default
        python3 Scraper.py -v  # verbose"""
    )

    arg_parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Verbose mode - output runtime details'
    )
    arg_parser.add_argument(
        '--debug-url',
        type=str,
        metavar='URL',
        help='Dry-run crawl: show discovered subpages and files for a given URL'
    )

    args = arg_parser.parse_args()

    if args.debug_url:
        debug_crawler(args.debug_url)
    else:
        build_vectorstore('sources.json')

